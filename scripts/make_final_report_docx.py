"""
Sinh báo cáo kết môn DOCX từ dự án Phân tích & Gợi ý BĐS.
Yêu cầu: python-docx, pandas, numpy
Chạy: python scripts/make_final_report_docx.py
Output: reports/BaoCao_KetMon_LapTrinhKHDL.docx
"""

import os, sys, json
from pathlib import Path
import pandas as pd
import numpy as np

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Cần cài: pip install python-docx")
    sys.exit(1)

# ── Đường dẫn ──────────────────────────────────────────────────────────────
BASE = Path(__file__).parent.parent.resolve()
REPORTS = BASE / "reports"
OUT = REPORTS / "BaoCao_KetMon_LapTrinhKHDL.docx"
os.makedirs(REPORTS, exist_ok=True)

# ── Màu sắc ────────────────────────────────────────────────────────────────
DARK_BLUE = RGBColor(0x1F, 0x49, 0x7D)   # Đậm
MID_BLUE  = RGBColor(0x2E, 0x74, 0xB5)   # Trung
LIGHT_BLUE = RGBColor(0xDE, 0xEB, 0xF7)   # Nền bảng
ORANGE     = RGBColor(0xC5, 0x50, 0x0C)   # Nhấn mạnh
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
BLACK      = RGBColor(0x00, 0x00, 0x00)
GRAY       = RGBColor(0xF2, 0xF2, 0xF2)

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Đặt màu nền ô bảng (hex không dấu #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.upper().lstrip("#"))
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    """Thiết lập đường viền ô."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:color"), "2E74B5")
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def para_fmt(para, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    pf = para.paragraph_format
    pf.alignment = align
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)


def heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    para_fmt(p, WD_ALIGN_PARAGRAPH.LEFT)
    if color:
        for run in p.runs:
            run.font.color.rgb = color
    return p


def add_paragraph(doc, text, bold=False, italic=False,
                  color=None, size=11, space=6, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    para_fmt(p, align, 0, space)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_figure(doc, rel_path_str, caption, width=Inches(5.5)):
    """Chèn hình ảnh với chú thích."""
    img_path = BASE / rel_path_str
    if img_path.exists():
        p = doc.add_paragraph()
        para_fmt(p, WD_ALIGN_PARAGRAPH.CENTER, 6, 4)
        run = p.add_run()
        run.add_picture(str(img_path), width=width)
    else:
        add_paragraph(doc, f"[Hình: {rel_path_str} — file không tìm thấy]", italic=True, color=RGBColor(0x80,0x80,0x80))
    cap = doc.add_paragraph(caption)
    para_fmt(cap, WD_ALIGN_PARAGRAPH.CENTER, 0, 12)
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    return cap


def add_table(doc, headers, rows, col_widths=None, header_bg="1F497D"):
    """Tạo bảng với header đậm."""
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = WHITE
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(hdr_cells[i], header_bg)
        set_cell_border(hdr_cells[i])

    # Data rows
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        bg = "DEEBF7" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)
            cells[c_idx].paragraphs[0].runs[0].font.size = Pt(9.5)
            if c_idx == 0:
                cells[c_idx].paragraphs[0].runs[0].bold = True
            set_cell_bg(cells[c_idx], bg)
            set_cell_border(cells[c_idx])

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    return table


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    para_fmt(p, WD_ALIGN_PARAGRAPH.LEFT, 0, 3)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def hr(doc):
    p = doc.add_paragraph()
    para_fmt(p, 0, 6, 6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "2E74B5")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_code_block(doc, code_text, caption=""):
    """Đoạn code giữ nguyên định dạng."""
    if caption:
        cap = doc.add_paragraph(caption)
        para_fmt(cap, WD_ALIGN_PARAGRAPH.LEFT, 6, 2)
        for run in cap.runs:
            run.bold = True
            run.font.size = Pt(10)
    p = doc.add_paragraph()
    para_fmt(p, WD_ALIGN_PARAGRAPH.LEFT, 0, 4)
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)


# ── Load data ────────────────────────────────────────────────────────────────

metrics = json.loads((BASE / "reports" / "metrics.json").read_text())
df_raw = pd.read_csv(BASE / "data" / "processed" / "listings_with_amenities.csv"
                      ).dropna(subset=["price_per_m2"])
df_cl = pd.read_csv(BASE / "data" / "processed" / "listings_with_clusters.csv")

# ── Main document ────────────────────────────────────────────────────────────

doc = Document()

# ── Page setup: A4, margins 2.5cm ──────────────────────────────────────────
from docx.oxml.ns import nsmap
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── Styles ────────────────────────────────────────────────────────────────────
styles = doc.styles
pStyle = styles["Normal"]
pStyle.font.name = "Times New Roman"
pStyle.font.size = Pt(13)

h1 = styles["Heading 1"]
h1.font.name = "Times New Roman"
h1.font.size = Pt(16)
h1.font.bold = True
h1.font.color.rgb = DARK_BLUE

h2 = styles["Heading 2"]
h2.font.name = "Times New Roman"
h2.font.size = Pt(14)
h2.font.bold = True
h2.font.color.rgb = MID_BLUE

h3 = styles["Heading 3"]
h3.font.name = "Times New Roman"
h3.font.size = Pt(12)
h3.font.bold = True
h3.font.color.rgb = DARK_BLUE

# =============================================================================
# TRANG BÌA
# =============================================================================
for _ in range(3):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run("BÁO CÁO KẾT MÔN")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = DARK_BLUE
r.font.name = "Times New Roman"

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub_p.add_run("Môn: Lập trình cho Khoa học Dữ liệu")
r2.font.size = Pt(14)
r2.font.name = "Times New Roman"

for _ in range(2):
    doc.add_paragraph()

topic_p = doc.add_paragraph()
topic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = topic_p.add_run("CHUYÊN ĐỀ 3")
r3.bold = True
r3.font.size = Pt(16)
r3.font.color.rgb = ORANGE
r3.font.name = "Times New Roman"

topic2_p = doc.add_paragraph()
topic2_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = topic2_p.add_run("PHÂN TÍCH VÀ GỢI Ý BẤT ĐỘNG SẢN\nCĂN HỘ / CHUNG CƯ TẠI TP. HỒ CHÍ MINH")
r4.bold = True
r4.font.size = Pt(18)
r4.font.color.rgb = DARK_BLUE
r4.font.name = "Times New Roman"

for _ in range(4):
    doc.add_paragraph()

info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
ri = info_p.add_run("Học viên: [Họ tên]\nMSSV: [MSSV]\nNgày nộp: 09/08/2026")
ri.font.size = Pt(12)
ri.font.name = "Times New Roman"

doc.add_page_break()

# =============================================================================
# MỤC LỤC (placeholder)
# =============================================================================
heading(doc, "MỤC LỤC", 1, DARK_BLUE)
toc_items = [
    "LỜI MỞ ĐẦU",
    "I.   Lý do lựa chọn đề tài",
    "II.  Mục đích nghiên cứu",
    "III. Phương pháp và phạm vi nghiên cứu",
    "IV.  Kết cấu đề tài",
    "Chương 1: Giới thiệu bài toán",
    "Chương 2: Tổng quan lý thuyết nền tảng",
    "Chương 3: Triển khai thuật toán",
    "Chương 4: Thực nghiệm và kết quả đạt được",
    "KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN",
    "TÀI LIỆU THAM KHẢO",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    para_fmt(p, WD_ALIGN_PARAGRAPH.LEFT, 0, 4)

doc.add_page_break()

# =============================================================================
# LỜI MỞ ĐẦU
# =============================================================================
heading(doc, "LỜI MỞ ĐẦU", 1, DARK_BLUE)

heading(doc, "I. Lý do lựa chọn đề tài", 2, MID_BLUE)
add_paragraph(doc,
    "Thị trường bất động sản (BĐS) căn hộ/chung cư tại TP. Hồ Chí Minh (TP.HCM) là một trong "
    "những thị trường sôi động nhất Việt Nam, với hàng nghìn tin đăng được đăng tải mỗi ngày "
    "trên các nền tảng rao bán như chotot.com, batdongsan.com.vn, alonhadat.com.vn. Tuy nhiên, "
    "dữ liệu từ các nguồn này thường không được chuẩn hóa: giá được viết dạng text (\"3,19 tỷ\", "
    "\"2,1 tỷ\"), hướng nhà và hướng ban công mã hóa bằng số (1–8), tình trạng nội thất và pháp "
    "lý cũng mã hóa bằng số (1–6). Điều này tạo ra nhu cầu thực tiễn cấp bách cho việc xây "
    "dựng một pipeline phân tích dữ liệu để chuẩn hóa, khám phá, dự đoán giá và gợi ý tin đăng "
    "phù hợp cho người dùng.", space=8)

add_paragraph(doc,
    "Bên cạnh đó, đồ án này còn là cơ hội để áp dụng toàn diện kiến thức từ môn Lập trình cho "
    "Khoa học Dữ liệu: thu thập và làm sạch dữ liệu, phân tích khám phá (EDA), xây dựng mô "
    "hình học máy có giám sát (supervised learning), phân cụm không giám sát (unsupervised "
    "learning) và hệ gợi ý (recommendation system). Đây là một trong những đồ án tổng hợp "
    "toàn diện nhất trong chương trình đào tạo.", space=8)

heading(doc, "II. Mục đích nghiên cứu", 2, MID_BLUE)
add_paragraph(doc,
    "Đồ án hướng đến các mục tiêu cụ thể sau:", space=6)
bullet(doc, "Thu thập và chuẩn hóa dữ liệu tin đăng BĐS căn hộ tại 6 quận TP.HCM từ nguồn dữ liệu tổng hợp (file xlsx snapshot).")
bullet(doc, "Phát hiện và xử lý các giá trị ngoại lệ (outlier) về diện tích, số phòng, giá bán.")
bullet(doc, "Xây dựng mô hình dự đoán giá/m² bằng các thuật toán học máy có giám sát (Linear Regression, Random Forest, Gradient Boosting).")
bullet(doc, "Phân cụm phân khúc thị trường căn hộ bằng thuật toán K-Means với việc tự động chọn số cụm K bằng silhouette score.")
bullet(doc, "Xây dựng hệ gợi ý top-5 tin đăng phù hợp với hồ sơ nhu cầu cụ thể của người dùng (ngân sách, quận, số phòng, diện tích).")
bullet(doc, "Trực quan hóa dữ liệu bằng các biểu đồ có tiêu đề, nhãn trục và nhận xét phân tích.")

heading(doc, "III. Phương pháp và phạm vi nghiên cứu", 2, MID_BLUE)
add_paragraph(doc,
    "Phương pháp nghiên cứu: Đồ án áp dụng quy trình Khoa học Dữ liệu (Data Science Lifecycle) "
    "gồm 9 bước theo chuẩn CRISP-DM: Xác định bài toán → Thu thập dữ liệu → Kiểm tra dữ liệu → "
    "Làm sạch → EDA → Xây dựng mô hình → Đánh giá → Triển khai → Báo cáo. Toàn bộ mã nguồn "
    "được viết bằng Python, sử dụng thư viện pandas, numpy, scikit-learn, matplotlib. Các mô hình "
    "học máy được đánh giá bằng 5-fold Cross-Validation trên tập train và đánh giá độc lập trên "
    "tập test (tỷ lệ 80/20).", space=8)

add_paragraph(doc,
    "Phạm vi nghiên cứu: Dữ liệu gồm 1799 tin đăng căn hộ/chung cư tại 6 quận TP.HCM "
    "(Thành phố Thủ Đức, Quận Bình Thạnh, Quận 7, Quận Gò Vấp, Quận 12, Quận Bình Tân), "
    "thu thập dưới dạng snapshot từ nền tảng chotot.com (tháng 07–08/2026). Nguồn dữ liệu phụ "
    "gồm 91 dòng thông tin tiện ích theo phường/xã (trường học, bệnh viện, chợ, công viên, "
    "bến xe buýt). Không phân tích xu hướng dài hạn do chỉ có dữ liệu snapshot.", space=8)

heading(doc, "IV. Kết cấu đề tài", 2, MID_BLUE)
add_paragraph(doc, "Đề tài được tổ chức thành 4 chương chính:", space=6)
bullet(doc, "Chương 1: Giới thiệu bài toán — Bối cảnh, mục tiêu, câu hỏi nghiên cứu và từ điển dữ liệu.")
bullet(doc, "Chương 2: Tổng quan lý thuyết nền tảng — Lý thuyết về các thuật toán học máy đã sử dụng (Linear Regression, Random Forest, Gradient Boosting, K-Means, Hệ gợi ý hybrid).")
bullet(doc, "Chương 3: Triển khai thuật toán — Kiến trúc hệ thống, thiết kế pipeline, các module mã nguồn, quy trình xử lý dữ liệu.")
bullet(doc, "Chương 4: Thực nghiệm và kết quả — Kết quả đánh giá mô hình, phân cụm, gợi ý, phân tích sai số và giới hạn.")

doc.add_page_break()

# =============================================================================
# CHƯƠNG 1: GIỚI THIỆU BÀI TOÁN
# =============================================================================
heading(doc, "Chương 1: GIỚI THIỆU BÀI TOÁN", 1, DARK_BLUE)

heading(doc, "1.1. Bối cảnh và thực trạng dữ liệu", 2, MID_BLUE)
add_paragraph(doc,
    "Dữ liệu tin đăng BĐS từ các nền tảng rao bán trực tuyến có những đặc thù riêng biệt "
    "khiến việc phân tích gặp nhiều thách thức. Cụ thể, giá bán thường được viết dưới dạng "
    "text có định dạng Việt Nam (\"3,19 tỷ\", \"2,1 tỷ\"), đòi hỏi phải parse chính xác sang "
    "dạng số. Hướng nhà và hướng ban công được mã hóa bằng các số nguyên 1 đến 8 (tương ứng "
    "8 hướng chính: Đông, Tây, Nam, Bắc, Đông Nam, Tây Nam, Đông Bắc, Tây Bắc). Tình trạng "
    "nội thất và pháp lý cũng mã hóa bằng các số nguyên không liên tục (1, 2, 4, 5, 6 cho pháp lý). "
    "Ngoài ra, dữ liệu còn chứa nhiều giá trị ngoại lệ rõ ràng: căn hộ có diện tích 1323 m² "
    "(gấp 16 lần trung bình), căn có 11 phòng ngủ (không hợp lý với loại hình căn hộ), và tỷ lệ "
    "dữ liệu thiếu (missing) cao: hướng nhà thiếu 79%, hướng ban công thiếu 73%, tình trạng "
    "nội thất thiếu 34%. Những thực trạng này đặt ra yêu cầu cấp thiết về việc xây dựng "
    "quy trình chuẩn hóa và làm sạch dữ liệu có hệ thống [1].", space=8)

heading(doc, "1.2. Mục tiêu và câu hỏi nghiên cứu", 2, MID_BLUE)
add_paragraph(doc, "Đồ án đặt ra 6 câu hỏi nghiên cứu cụ thể:", space=6)
bullet(doc, "Trong 6 quận có dữ liệu, quận nào có giá/m² cao nhất và thấp nhất?")
bullet(doc, "Diện tích và số phòng ngủ ảnh hưởng thế nào đến giá/m²?")
bullet(doc, "Hướng nhà, hướng ban công, tình trạng nội thất và pháp lý có mối liên hệ gì với giá?")
bullet(doc, "Mô hình dự đoán giá/m² đạt sai số bao nhiêu trên thực tế (MAE, RMSE, R²)?")
bullet(doc, "Thị trường căn hộ TP.HCM có thể phân thành bao nhiêu phân khúc theo đặc điểm?")
bullet(doc, "Hệ gợi ý có thể đề xuất top-5 tin đăng phù hợp với hồ sơ nhu cầu cụ thể không?")

heading(doc, "1.3. Từ điển dữ liệu", 2, MID_BLUE)
add_paragraph(doc,
    "Dữ liệu chính gồm 1799 tin đăng căn hộ/chung cư từ file xlsx với 32 cột gốc, "
    "sau khi chuẩn hóa và bổ sung cột tính toán có tổng cộng 41 cột. Bảng dưới đây trình "
    "bày các thuộc tính chính được sử dụng trong phân tích và xây dựng mô hình:", space=8)

add_table(doc,
    ["Thuộc tính", "Kiểu", "Missing (%)", "Vai trò", "Ghi chú"],
    [
        ["listing_id", "int", "0.0%", "Khóa chính", "Mã tin đăng duy nhất"],
        ["total_price", "int", "0.0%", "Target phụ", "Tổng giá bán (VND)"],
        ["area_m2", "float", "0.3%", "Feature chính", "Diện tích căn hộ (m²)"],
        ["price_per_m2", "float", "0.3%", "Target chính", "Giá/m² (VND), tính lại từ total_price/area_m2"],
        ["bedrooms", "float", "0.2%", "Feature chính", "Số phòng ngủ"],
        ["bathrooms", "float", "15.6%", "Feature phụ", "Số phòng vệ sinh"],
        ["direction", "float (1–8)", "78.9%", "Feature", "Mã hướng nhà → decode sang tên"],
        ["balcony_direction", "float (1–8)", "73.2%", "Feature", "Mã hướng ban công"],
        ["furnishing_status", "float (1–4)", "34.4%", "Feature", "Mã nội thất (1=Không, 2=Cơ bản, 3=Đầy đủ, 4=Cao cấp)"],
        ["legal_status", "float (1,2,4,5,6)", "20.0%", "Feature", "Mã pháp lý"],
        ["apartment_type", "int (1–6)", "0.0%", "Feature", "Loại căn hộ (1=phổ biến nhất)"],
        ["image_count", "int", "0.0%", "Feature phụ", "Số ảnh — chất lượng tin đăng"],
        ["project_name", "str", "20.9%", "Feature categorical", "Tên dự án → OHE với min_freq=10"],
        ["district_clean", "str", "0.0%", "Feature categorical", "Quận/huyện đã chuẩn hóa (6 quận)"],
        ["amenity_score", "float", "0.0%", "Feature phụ", "Điểm tiện ích theo phường (từ nguồn phụ)"],
        ["cluster", "int (0–3)", "0.0%", "Phân cụm", "Kết quả K-Means, K=4"],
    ],
    col_widths=[1.8, 1.5, 0.9, 1.2, 2.6]
)

heading(doc, "1.4. Quy trình làm sạch dữ liệu", 2, MID_BLUE)
add_paragraph(doc,
    "Quy trình làm sạch gồm các bước: (1) Parse giá text → số, (2) Chuẩn hóa tên quận, "
    "(3) Decode mã hướng nhà 1–8, (4) Decode mã nội thất và pháp lý, (5) Lọc outlier: "
    "diện tích < 10 hoặc > 500 m² (loại 3 dòng), số phòng ngủ > 10 (loại 17 dòng), "
    "giá/m² < 100 triệu hoặc > 500 triệu. Kết quả: 1799 dòng raw → 1779 dòng sau "
    "outlier filter → 1773 dòng sau dropna(subset=['price_per_m2']). Chi tiết được ghi "
    "trong cleaning_log.csv và error_log.txt [2].", space=8)

add_figure(doc, "reports/figures/fig05_listings_by_district.png",
    "Hình 1.1. Phân bố số tin đăng theo 6 quận TP.HCM trong tập dữ liệu sau làm sạch (1773 dòng).")

doc.add_page_break()

# =============================================================================
# CHƯƠNG 2: TỔNG QUAN LÝ THUYẾT NỀN TẢNG
# =============================================================================
heading(doc, "Chương 2: TỔNG QUAN LÝ THUYẾT NỀN TẢNG", 1, DARK_BLUE)

heading(doc, "2.1. Mô hình hồi quy tuyến tính (Linear Regression)", 2, MID_BLUE)
add_paragraph(doc,
    "Hồi quy tuyến tính là mô hình cơ bản nhất trong học máy có giám sát, mục tiêu tìm "
    "một hàm tuyến tính f(x) = w·x + b sao cho tổng bình phương sai số ∑(y_i − f(x_i))² "
    "là nhỏ nhất. Trong bài toán này, ta sử dụng biến đổi log1p lên target (giá/m²) trước "
    "khi fit: y_transformed = log(1 + price_per_m2). Mục đích là giảm skewness của phân bố "
    "giá (lệch phải nặng), giúp residual gần phân bố chuẩn hơn và MAE/RMSE ổn định hơn [3]. "
    "Sau khi predict, kết quả được chuyển ngược lại bằng expm1(y). Đây là phương pháp "
    "phổ biến trong dự đoán giá BĐS vì giá BĐS có phân bố log-normal.", space=8)

heading(doc, "2.2. Random Forest Regressor", 2, MID_BLUE)
add_paragraph(doc,
    "Random Forest là một phương pháp ensemble learning, kết hợp dự đoán của nhiều cây "
    "quyết định (decision trees) để giảm overfitting và tăng độ chính xác tổng thể. Mỗi cây "
    "được huấn luyện trên một tập con của dữ liệu (bootstrap sampling) và chỉ xem xét "
    "một tập con ngẫu nhiên của các features tại mỗi node — đây là hai nguồn \"ngẫu nhiên\" "
    "tạo nên tên gọi \"Random\" Forest.", space=8)
add_paragraph(doc,
    "Cụ thể trong đồ án này, thuật toán được cấu hình với các siêu tham số: "
    "n_estimators=200 (200 cây), max_depth=None (cây phát triển đầy đủ), "
    "min_samples_leaf=2 (tối thiểu 2 mẫu tại mỗi lá để tránh overfitting), "
    "n_jobs=-1 (song song hóa trên tất cả CPU cores). Kết quả dự đoán của toàn bộ "
    "cây được tính trung bình (averaging) để cho ra dự đoán cuối cùng. Random Forest "
    "nổi bật ở khả năng học các quan hệ phi tuyến và tương tác giữa các features "
    "(ví dụ: area_m2 × district_clean) mà không cần feature engineering thủ công [4].", space=8)
add_paragraph(doc,
    "Điểm mạnh của Random Forest so với Hồi quy tuyến tính trong bài toán BĐS: "
    "(a) Không giả định phân phối chuẩn của residuals, (b) Tự động nắm bắt tương tác "
    "phi tuyến giữa location và property features, (c) Ít nhạy cảm với outlier hơn "
    "Linear Regression (cây quyết định split theo threshold, không bị ảnh hưởng bởi "
    "giá trị tuyệt đối của outlier), (d) Cung cấp feature importance tự nhiên — "
    "đo mức đóng góp của mỗi feature trong việc giảm RSS qua tất cả các cây.", space=8)

heading(doc, "2.3. Gradient Boosting Regressor", 2, MID_BLUE)
add_paragraph(doc,
    "Gradient Boosting là kỹ thuật ensemble learning tuần tự (sequential), khác với "
    "Random Forest là parallel. Thuật toán xây dựng các cây quyết định một cách tuần tự, "
    "mỗi cây mới được huấn luyện để sửa sai số (residual) của tất cả các cây trước đó "
    "đó. Cụ thể, tại mỗi vòng lặp, thuật toán tính gradient của hàm mất mát (loss "
    "function) theo dự đoán hiện tại, rồi fit một cây mới để dự đoán gradient đó "
    "(stage-wise). Hàm mất mát mặc định là squared error: L = (y − ŷ)².", space=8)
add_paragraph(doc,
    "Trong đồ án này, Gradient Boosting được cấu hình: n_estimators=200 (200 stages), "
    "max_depth=4 (cây nông, chỉ 4 mức split — tránh overfitting trên dataset nhỏ), "
    "learning_rate=0.05 (tốc độ học thấp để tăng số lượng cây bù lại, cải thiện "
    "generalization). Gradient Boosting thường đạt hiệu suất cao hơn Random Forest "
    "trên các tập dữ liệu có tương tác phức tạp, nhưng nhạy cảm hơn với overfitting "
    "khi n_estimators quá lớn hoặc max_depth quá sâu [5].", space=8)

heading(doc, "2.4. K-Means Clustering", 2, MID_BLUE)
add_paragraph(doc,
    "K-Means là thuật toán phân cụm phổ biến nhất trong học máy không giám sát "
    "(unsupervised learning). Mục tiêu là phân n dữ liệu điểm vào K cụm sao cho "
    "tổng bình phương khoảng cách từ mỗi điểm đến tâm cụm (centroid) gần nhất "
    "(within-cluster sum of squares — WCSS) là nhỏ nhất. Thuật toán gồm 2 bước lặp: "
    "(1) E-step: gán mỗi điểm cho cụm có centroid gần nhất (theo khoảng cách Euclidean), "
    "(2) M-step: tính lại centroid cho mỗi cụm = trung bình các điểm thuộc cụm. "
    "Hội tụ khi các phép gán không thay đổi hoặc số vòng lặp đạt giới hạn.", space=8)
add_paragraph(doc,
    "Trước khi áp dụng K-Means, dữ liệu được chuẩn hóa bằng StandardScaler "
    "(trừ trung bình, chia độ lệch chuẩn) vì K-Means sử dụng khoảng cách "
    "Euclidean — các feature có scale khác nhau sẽ gây sai lệch. Việc chọn K "
    "được thực hiện tự động bằng Silhouette Score: s = (b − a) / max(a, b), "
    "trong đó a là khoảng cách trung bình trong cụm, b là khoảng cách trung bình "
    "đến cụm gần nhất tiếp theo. Silhouette ∈ [−1, 1]: gần 1 nghĩa là cụm tách biệt "
    "rõ ràng, gần 0 nghĩa là các cụm chồng lấn, âm nghĩa là có điểm bị gán nhầm cụm. "
    "K có Silhouette cao nhất được chọn [6].", space=8)

heading(doc, "2.5. Hệ gợi ý Hybrid (Hybrid Recommendation)", 2, MID_BLUE)
add_paragraph(doc,
    "Hệ gợi ý hybrid trong đồ án này kết hợp hai chiến lược: (a) Lọc cứng "
    "(content-based filtering) dựa trên các ràng buộc về ngân sách, quận và số phòng "
    "ngủ; (b) Chấm điểm có trọng số (weighted scoring) dựa trên mức độ phù hợp về "
    "giá, diện tích, cùng cụm K-Means và điểm tiện ích xung quanh.", space=8)
add_paragraph(doc,
    "Công thức tính điểm tổng hợp cho mỗi tin đăng ứng viên:", space=6)
add_code_block(doc,
"score_total = price_score + area_score + segment_bonus + amenity_bonus\n"
"\n"
"price_score   = 1 − |price_per_m2 − target_price_per_m2| / target_price_per_m2\n"
"area_score    = 1 − |area_m2 − target_area| / target_area\n"
"segment_bonus = 0.3  nếu cluster == preferred_cluster,  ngược lại 0\n"
"amenity_bonus = 0.2 × (amenity_score / max_amenity_score)\n"
"\n"
"trong đó target_price_per_m2 = budget / target_area",
"   Công thức tính điểm trong hệ gợi ý hybrid.")
add_paragraph(doc,
    "Ưu điểm của chiến lược hybrid so với collaborative filtering truyền thống: "
    "(a) Không cần dữ liệu lịch sử hành vi người dùng (cold-start problem), "
    "(b) Giải thích được tại sao một tin được gợi ý thông qua từng thành phần điểm "
    "(score components), (c) Linh hoạt với các hồ sơ nhu cầu mới. Đây là phương pháp "
    "phổ biến trong các ứng dụng BĐS khi chưa có đủ dữ liệu hành vi người dùng [7].", space=8)

doc.add_page_break()

# =============================================================================
# CHƯƠNG 3: TRIỂN KHAI THUẬT TOÁN
# =============================================================================
heading(doc, "Chương 3: TRIỂN KHAI THUẬT TOÁN", 1, DARK_BLUE)

heading(doc, "3.1. Kiến trúc tổng thể hệ thống", 2, MID_BLUE)
add_paragraph(doc,
    "Hệ thống được tổ chức theo kiến trúc pipeline xử lý dữ liệu tuần tự (data pipeline), "
    "gồm 4 giai đoạn chính. Giai đoạn 1 (Thu thập & Làm sạch): file xlsx 1799 dòng "
    "→ cleaner.py (chuẩn hóa district, decode direction/furnishing/legal) → "
    "filter outlier (area, bedrooms, price) → listings_clean.csv (1779 dòng). "
    "Giai đoạn 2 (Bổ sung & Transform): merge amenity_score theo (district, ward) → "
    "build_preprocessor (ColumnTransformer: impute + scale + OHE) → 47 features "
    "transformed. Giai đoạn 3 (Mô hình học máy): 4 mô hình (Dummy/Linear/RF/GBR) "
    "→ 5-fold CV → đánh giá test set → K-Means (silhouette auto-pick K). "
    "Giai đoạn 4 (Gợi ý): RecommendationEngine cho 3 hồ sơ mẫu → top-5.", space=8)

add_figure(doc, "reports/figures/fig06_correlation_heatmap.png",
    "Hình 3.1. Heatmap tương quan giữa các biến số trong tập dữ liệu căn hộ. "
    "Hệ số tương quan Pearson được tính cho các thuộc tính numeric. "
    "Nhận xét: tương quan giữa price_per_m2 và các biến khác yếu (<0.3), "
    "bedrooms ↔ area_m2 = 0.55 (tương quan dự kiến), "
    "total_price ↔ area_m2 = 0.70.")

heading(doc, "3.2. Thiết kế Feature Pipeline", 2, MID_BLUE)
add_paragraph(doc,
    "Feature pipeline được xây dựng bằng sklearn.compose.ColumnTransformer với 2 nhánh "
    "xử lý song song:", space=6)
bullet(doc, "Nhánh numeric (9 features): SimpleImputer(strategy='median') → StandardScaler(). "
           "Median imputation được chọn thay vì mean vì ít bị ảnh hưởng bởi outlier. "
           "StandardScaler đảm bảo các features có cùng scale (mean=0, std=1) — "
           "điều kiện tiên quyết cho K-Means (khoảng cách Euclidean).")
bullet(doc, "Nhánh categorical (3 features): SimpleImputer(strategy='constant', fill_value='missing') "
           "→ OneHotEncoder(handle_unknown='ignore', min_frequency=10). "
           "handle_unknown='ignore' đảm bảo pipeline robust khi gặp giá trị mới ở test. "
           "min_frequency=10 gộp các category hiếm thành một dummy variable → giảm chiều.")

add_paragraph(doc,
    "Tổng cộng sau transform: 9 numeric + 38 categorical OHE = **47 features**. "
    "Pipeline này được fit trên train set (1418 dòng) và transform độc lập trên "
    "test set (355 dòng) — không có data leakage.", space=8)

add_code_block(doc,
"# src/features.py — ColumnTransformer pipeline\n"
"numeric_pipe = Pipeline([\n"
"    ('impute', SimpleImputer(strategy='median')),\n"
"    ('scale',  StandardScaler()),\n"
"])\n"
"categorical_pipe = Pipeline([\n"
"    ('impute', SimpleImputer(strategy='constant', fill_value='missing')),\n"
"    ('ohe',    OneHotEncoder(handle_unknown='ignore', min_frequency=10)),\n"
"])\n"
"ct = ColumnTransformer([\n"
"    ('num', numeric_pipe,    NUMERIC_COLS),   # 9 cols\n"
"    ('cat', categorical_pipe, CATEGORICAL_COLS), # 3 cols → 38 dummies\n"
"], remainder='drop', sparse_threshold=0.0)\n"
"# Fit trên train, transform cả train và test",
"[src/features.py]")

heading(doc, "3.3. Thiết kế mô hình PricePredictor", 2, MID_BLUE)
add_paragraph(doc,
    "Module src/predictor.py đóng gói 4 mô hình vào class PricePredictor có interface "
    "thống nhất: fit(X, y) → predict(X) → evaluate(X_test, y_test). "
    "Hàm cv_metrics() thực hiện 5-fold Cross-Validation trên toàn bộ train set (1418 dòng):", space=6)
add_code_block(doc,
"# src/predictor.py — 4 mô hình\n"
"def _build_model(model_name, random_state):\n"
"    if model_name == 'dummy':\n"
"        return DummyRegressor(strategy='median')\n"
"    if model_name == 'linear':\n"
"        return LinearRegression()\n"
"    if model_name == 'rf':   # Random Forest\n"
"        return RandomForestRegressor(\n"
"            n_estimators=200, max_depth=None,\n"
"            min_samples_leaf=2, n_jobs=-1,\n"
"            random_state=random_state)\n"
"    if model_name == 'gbr':  # Gradient Boosting\n"
"        return GradientBoostingRegressor(\n"
"            n_estimators=200, max_depth=4,\n"
"            learning_rate=0.05,\n"
"            random_state=random_state)\n"
"\n"
"def cv_metrics(X, y, model_name, log_target=True, n_splits=5):\n"
"    # KFold(5, shuffle=True, random_state=42)\n"
"    # Trả về mae_mean±std, rmse_mean±std, r2_mean±std",
"[src/predictor.py]")

heading(doc, "3.4. Thiết kế KMeansSegmenter với Silhouette Auto-pick K", 2, MID_BLUE)
add_paragraph(doc,
    "Module src/segmenter.py gói Pipeline(StandardScaler + KMeans) trong class "
    "KMeansSegmenter. Hàm pick_k_by_silhouette() thử lần lượt K từ 3 đến 6, "
    "tính silhouette_score trên toàn bộ train set, và chọn K có silhouette cao nhất. "
    "Đây là phương pháp chọn K không cần labeled data (phù hợp với unsupervised learning).", space=8)

add_code_block(doc,
"# src/segmenter.py — Auto-pick K bằng silhouette\n"
"def pick_k_by_silhouette(X, k_range=(3, 7), random_state=42):\n"
"    scores = {}\n"
"    for k in range(k_range[0], k_range[1] + 1):\n"
"        pipe = Pipeline([\n"
"            ('scale', StandardScaler()),\n"
"            ('km',   KMeans(n_clusters=k, n_init=10,\n"
"                           random_state=random_state))\n"
"        ])\n"
"        labels = pipe.fit_predict(X)\n"
"        s = silhouette_score(X, labels)\n"
"        scores[k] = float(s)\n"
"    best_k = max(scores, key=scores.get)\n"
"    return best_k, scores",
"[src/segmenter.py]")

heading(doc, "3.5. Thiết kế RecommendationEngine", 2, MID_BLUE)
add_paragraph(doc,
    "Module src/recommender.py cài đặt hệ gợi ý hybrid theo 3 bước: "
    "(1) Filter cứng: tổng giá ∈ [0.8×budget, 1.2×budget], số phòng ∈ [target−1, target+1], "
    "district ∈ preferred_districts. (2) Scoring: tính price_score, area_score, "
    "segment_bonus (nếu cùng cluster K-Means), amenity_bonus (tỷ lệ amenity_score). "
    "(3) Sắp xếp theo score_total giảm dần, lấy top_k=5.", space=8)
add_paragraph(doc,
    "Thiết kế này cho phép người dùng tùy biến hồ sơ nhu cầu (budget, bedrooms, "
    "target_area, preferred_districts, preferred_cluster) mà không cần thay đổi code. "
    "Score components được lưu lại để giải thích từng thành phần điểm — quan trọng "
    "cho việc xây dựng trust với người dùng cuối.", space=8)

doc.add_page_break()

# =============================================================================
# CHƯƠNG 4: THỰC NGHIỆM VÀ KẾT QUẢ
# =============================================================================
heading(doc, "Chương 4: THỰC NGHIỆM VÀ KẾT QUẢ ĐẠT ĐƯỢC", 1, DARK_BLUE)

heading(doc, "4.1. Môi trường thực nghiệm và dữ liệu", 2, MID_BLUE)
add_table(doc,
    ["Thành phần", "Mô tả"],
    [
        ["Ngôn ngữ", "Python 3.12+"],
        ["Thư viện chính", "pandas 2.0+, numpy 1.24+, scikit-learn 1.3+, matplotlib 3.7+"],
        ["Train / Test split", "80% / 20% (random_state=42)"],
        ["Số mẫu train", "1418 dòng"],
        ["Số mẫu test", "355 dòng"],
        ["Số features (sau transform)", "47 (9 numeric + 38 OHE)"],
        ["Cross-validation", "5-fold KFold (shuffle=True, random_state=42)"],
        ["Random state", "42 (toàn bộ pipeline)"],
    ],
    col_widths=[2.5, 5.5]
)

heading(doc, "4.2. Kết quả so sánh các mô hình dự đoán giá/m²", 2, MID_BLUE)
add_paragraph(doc,
    "Bảng dưới đây tổng hợp kết quả đánh giá 4 mô hình trên cả tập "
    "Cross-Validation (5-fold, mean ± std) và tập test độc lập (355 mẫu). "
    "Target là price_per_m2 (VND), tất cả mô hình đều fit trên log1p(price_per_m2).", space=8)

rf = metrics["models"]["rf"]
lr = metrics["models"]["linear"]
gbr = metrics["models"]["gbr"]
dm = metrics["models"]["dummy"]

add_table(doc,
    ["Mô hình", "CV MAE (±std)", "CV RMSE (±std)", "CV R² (±std)",
     "Test MAE", "Test RMSE", "Test R²"],
    [
        ["Dummy (Baseline)", f"{dm['mae_mean']/1e6:.1f}M", f"{dm['rmse_mean']/1e6:.1f}M",
         f"{dm['r2_mean']:.3f}", f"{dm['test_mae']/1e6:.1f}M",
         f"{dm['test_rmse']/1e6:.1f}M", f"{dm['test_r2']:.3f}"],
        ["Linear Regression", f"{lr['mae_mean']/1e6:.1f}M", f"{lr['rmse_mean']/1e6:.1f}M",
         f"{lr['r2_mean']:.3f}", f"{lr['test_mae']/1e6:.1f}M",
         f"{lr['test_rmse']/1e6:.1f}M", f"{lr['test_r2']:.3f}"],
        ["Random Forest ★", f"{rf['mae_mean']/1e6:.1f}M", f"{rf['rmse_mean']/1e6:.1f}M",
         f"{rf['r2_mean']:.3f}", f"{rf['test_mae']/1e6:.1f}M",
         f"{rf['test_rmse']/1e6:.1f}M", f"{rf['test_r2']:.3f}"],
        ["Gradient Boosting", f"{gbr['mae_mean']/1e6:.1f}M", f"{gbr['rmse_mean']/1e6:.1f}M",
         f"{gbr['r2_mean']:.3f}", f"{gbr['test_mae']/1e6:.1f}M",
         f"{gbr['test_rmse']/1e6:.1f}M", f"{gbr['test_r2']:.3f}"],
    ],
    col_widths=[1.8, 1.2, 1.3, 1.1, 1.0, 1.1, 0.9]
)

add_paragraph(doc,
    "★ Mô hình tốt nhất trên tập test: **Random Forest** với Test R² = 0.169, "
    "Test MAE ≈ 13.9 triệu VND/m², Test RMSE ≈ 33.2 triệu VND/m². "
    "Sai số MAE tương đương khoảng 26% so với median giá/m² toàn mẫu (~53 triệu). "
    "Nhận xét: Dummy R² âm (−0.041) đúng kỳ vọng — median không dự đoán được gì. "
    "Linear R² = 0.143 — giải thích ~14% phương sai. RF vượt Linear (+0.026 R²) "
    "chứng minh có quan hệ phi tuyến. GBR CV R² = 0.366 nhỉnh hơn RF (CV) nhưng "
    "gap CV-Test = 0.204 lớn hơn RF (gap = 0.188) → RF generalize tốt hơn trên tập test [8].", space=8)

add_figure(doc, "reports/figures/fig04_area_vs_price.png",
    "Hình 4.1. Scatter plot thể hiện quan hệ giữa diện tích (m²) và tổng giá (VND). "
    "Quan hệ gần tuyến tính — diện tích càng lớn giá càng cao. "
    "Một số căn diện tích 50–80m² ở mức giá 3–5 tỷ → căn cao cấp vị trí đắt đỏ.")

heading(doc, "4.3. Phân tích 10 trường hợp dự đoán sai lớn nhất", 2, MID_BLUE)
add_paragraph(doc,
    "Bảng dưới liệt kê 10 trường hợp có sai số tuyệt đối (abs_error) lớn nhất "
    "trên tập test (355 mẫu), cùng với % sai số (pct_error) và listing_id tương ứng. "
    "pct_error = |actual − predicted| / actual × 100%.", space=8)

add_table(doc,
    ["STT", "listing_id", "Quận", "Diện tích (m²)", "PN", "Giá thực (tr/m²)", "Giá dự đoán (tr/m²)", "%Sai"],
    [
        ["1", "177641859", "Quận 12", "26.0", "1", "480.8", "42.2", "91.2%"],
        ["2", "177622744", "Q. Bình Thạnh", "60.0", "3", "258.3", "57.5", "77.8%"],
        ["3", "177494566", "Q. Bình Thạnh", "122.4", "3", "228.8", "72.4", "68.4%"],
        ["4", "177850739", "TP. Thủ Đức", "140.0", "3", "51.4", "127.3", "147.4%"],
        ["5", "177757531", "TP. Thủ Đức", "74.0", "2", "162.2", "58.9", "63.7%"],
        ["6", "177715817", "Quận 7", "82.0", "2", "146.3", "54.8", "62.5%"],
        ["7", "177770338", "Q. Bình Thạnh", "65.0", "2", "121.5", "55.9", "54.0%"],
        ["8", "177866544", "TP. Thủ Đức", "51.0", "1", "117.6", "52.4", "55.4%"],
        ["9", "177766301", "TP. Thủ Đức", "71.0", "2", "133.8", "70.2", "47.5%"],
        ["10", "175165640", "TP. Thủ Đức", "64.0", "1", "215.6", "98.5", "54.3%"],
    ],
    col_widths=[0.4, 1.1, 1.3, 1.1, 0.4, 1.3, 1.4, 0.8]
)

add_paragraph(doc,
    "Nhận xét: (1) **5/10 trường hợp** là căn có giá cực cao (> 150 triệu/m²) — "
    "các dự án cao cấp (Landmark 81, Empire City, Masteri An Phú, Midtown, Lumière Riverside). "
    "Mô hình dự đoán thấp hơn 30–60% vì giá thực phụ thuộc nhiều vào view sông, "
    "tầng cao, tiến độ dự án — không có trong features. "
    "(2) **1 trường hợp (listing 177850739)** có pct_error = 147.4%: căn Duplex thông tầng "
    "140m² ở Thủ Đức, actual = 51.4 triệu/m² bất thường thấp hơn nhiều so với "
    "model dự đoán (127.3M) — có thể là data entry error trong tin đăng gốc. "
    "(3) **5/10 trường hợp** thiếu từ 3–5 features trong số 9 features numeric → "
    "median imputation gây sai lệch cho 50% top-10 worst. "
    "Hướng cải thiện: thu thập thêm ≥ 10.000 tin, thêm feature tầng/view/khoảng cách "
    "đến trung tâm, thử XGBoost/LightGBM với hyperparameter tuning (Optuna) [9].", space=8)

heading(doc, "4.4. Kết quả phân cụm K-Means", 2, MID_BLUE)

# Silhouette table
km = metrics["kmeans"]
add_table(doc,
    ["K", "Silhouette Score", "Chọn"],
    [
        ["3", f"{km['scores']['3']:.3f}", "—"],
        ["4", f"{km['scores']['4']:.3f}", "✓ (tốt nhất)"],
        ["5", f"{km['scores']['5']:.3f}", "—"],
        ["6", f"{km['scores']['6']:.3f}", "—"],
    ],
    col_widths=[1.5, 2.5, 2.0]
)

add_paragraph(doc,
    "K = 4 được chọn tự động với Silhouette = 0.083. Đây là giá trị thấp (lý "
    "tưởng > 0.5), cho thấy các cụm không tách biệt rõ ràng trong không gian "
    "47 chiều — kỳ vọng với dữ liệu BĐS vì giá phân bố liên tục.", space=8)

# Cluster distribution table
add_table(doc,
    ["Cluster", "Số tin", "Tỷ lệ", "Median giá/m²", "Đặc điểm nổi bật"],
    [
        ["0", "165", "9.3%", "50.0M", "Căn 1PN, hướng Đông (direction_code +0.57), balcony Đông (+1.04)"],
        ["1", "1396", "78.7%", "53.7M", "Đa số — căn 2PN đặc trưng, phân bố trung bình"],
        ["2", "47", "2.6%", "50.0M", "Căn nhỏ ở Thủ Đức, scaled area = -0.45 (≈55m²)"],
        ["3", "165", "9.3%", "53.8M", "Hướng Tây/Tây Bắc (direction_code -2.27), balcony Tây (-1.70)"],
    ],
    col_widths=[0.8, 0.9, 0.8, 1.2, 3.3]
)

add_paragraph(doc,
    "Quan trọng: **4 cụm có median giá/m² gần như nhau (50–54M)** → K-Means "
    "phân biệt chủ yếu theo `direction_code` (hướng nhà), KHÔNG theo phân khúc giá. "
    "Đây là hệ quả trực tiếp của R² thấp (0.169) — các features không giải thích "
    "được nhiều phương sai của giá, nên K-Means cũng không thể phân tách theo giá. "
    "Silhouette thấp (0.083) xác nhận các cụm chồng lấn nhau.", space=8)

add_figure(doc, "reports/figures/fig11_silhouette.png",
    "Hình 4.2. Biểu đồ Silhouette score theo số cụm K. "
    "K=4 đạt silhouette cao nhất (0.083), K=6 có silhouette âm (−0.061) → cụm không ổn định. "
    "Đường màu cam đánh dấu giá trị silhouette trung bình cho mỗi K.")

heading(doc, "4.5. Kết quả hệ gợi ý top-5", 2, MID_BLUE)
add_paragraph(doc,
    "Hệ gợi ý được demo với 3 hồ sơ nhu cầu căn hộ tại TP.HCM. "
    "Chi tiết từng dòng trong file reports/sample_recommendations.csv.", space=8)

add_table(doc,
    ["Hồ sơ", "Ngân sách", "Phòng", "Diện tích (m²)", "Quận ưu tiên", "Cluster ưu tiên", "Kết quả"],
    [
        ["Gia đình trẻ", "3 tỷ", "2PN", "65", "Thủ Đức, Bình Thạnh", "0", "Top-5 ✓"],
        ["Nhà đầu tư", "5 tỷ", "2PN", "70", "Quận 7, Bình Tân", "1", "Top-5 ✓"],
        ["Người mua cao cấp", "7 tỷ", "3PN", "85", "Thủ Đức, Quận 7", "1*", "Top-5 ✓"],
    ],
    col_widths=[1.8, 1.0, 0.7, 1.0, 1.8, 1.3, 0.9]
)

add_paragraph(doc,
    "* Cluster ưu tiên ban đầu là 2 (chỉ 47 tin Thủ Đức, giá 1.58–4.6 tỷ → 0 tin "
    "khớp filter cứng cho profile 7 tỷ) → đã đổi sang cluster 1 (1396 tin, "
    "có nhiều căn 6–10 tỷ). Chú ý: `preferred_cluster` chỉ là bonus điểm (+0.3), "
    "KHÔNG phải filter cứng. District name phải khớp chính xác với data "
    "(Thành phố Thủ Đức, Quận 7, Quận Bình Thạnh, Quận Bình Tân, Quận 12, Quận Gò Vấp).", space=8)

add_figure(doc, "reports/figures/fig09_amenity_by_district.png",
    "Hình 4.3. Điểm tiện ích trung bình theo 6 quận (amenity_score). "
    "Thành phố Thủ Đức và Quận 7 có điểm cao nhất (~7.5–8.5) — nhiều dự án, "
    "tiện ích hiện đại. Bình Tân và Quận 12 thấp hơn (~5–6).")

doc.add_page_break()

# =============================================================================
# KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN
# =============================================================================
heading(doc, "KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", 1, DARK_BLUE)

heading(doc, "5.1. Kết luận", 2, MID_BLUE)
add_paragraph(doc,
    "Đồ án đã hoàn thành đầy đủ yêu cầu của Chuyên đề 3 với các kết quả chính sau:", space=8)
bullet(doc, "Xây dựng pipeline hoàn chỉnh gồm 9 bước CRISP-DM, xử lý 1799 tin đăng từ 2 nguồn dữ liệu (xlsx + csv) qua các giai đoạn: thu thập → kiểm tra → làm sạch → EDA → mô hình → đánh giá → báo cáo.")
bullet(doc, "Chuẩn hóa dữ liệu: decode 8 hướng nhà, 4 mức nội thất, 5 mức pháp lý; lọc 20 outlier (3 area + 17 bedrooms); ghi log chi tiết vào cleaning_log.csv và error_log.txt. Kết quả: 1773 dòng sạch.")
bullet(doc, "10 biểu đồ EDA có tiêu đề, nhãn trục, nhận xét phân tích. Phát hiện: Quận Bình Thạnh có median giá/m² cao nhất (71.4M), Bình Tân thấp nhất (42.8M); 2PN chiếm 67% (1189/1773 tin).")
bullet(doc, "So sánh 4 mô hình: Dummy (baseline), Linear Regression, Random Forest, Gradient Boosting. Random Forest đạt Test R² = 0.169, MAE = 13.9M/m² (~26% sai số so với median 53M). GBR CV R² = 0.366 nhỉnh hơn nhưng overfitting cao hơn.")
bullet(doc, "K-Means tự động chọn K=4 bằng Silhouette (0.083). 4 cụm khác biệt chủ yếu theo direction_code, KHÔNG theo phân khúc giá (median 4 cluster 50–54M, gần như nhau).")
bullet(doc, "Hệ gợi ý hybrid cho 3 hồ sơ nhu cầu: gia đình trẻ (3 tỷ Thủ Đức), nhà đầu tư (5 tỷ Quận 7), người mua cao cấp (7 tỷ Thủ Đức/Quận 7) — cả 3 profile đều có top-5.")
bullet(doc, "45/45 unit test PASS, 4 Notebook Jupyter chạy end-to-end, báo cáo Markdown + slide outline hoàn chỉnh.")

heading(doc, "5.2. Hạn chế và rủi ro", 2, MID_BLUE)
bullet(doc, "Chỉ 1773 dòng từ 1 nguồn snapshot T7–T8/2026 — không phân tích xu hướng dài hạn.")
bullet(doc, "Phạm vi 6 quận (thiếu Quận 1, 3, 4, 5) → bias về khu vực, thiếu phân khúc cao cấp trung tâm.")
bullet(doc, "R² = 0.17 — giá BĐS phụ thuộc nhiều yếu tố phi số (view sông, tầng cao, nội thất chi tiết, tiến độ dự án) không có trong data.")
bullet(doc, "Silhouette = 0.083 — K-Means khó phân tách rõ phân khúc thị trường.")
bullet(doc, "Missing cao: direction 79%, balcony 73%, furnishing 34% → median imputation gây bias.")
bullet(doc, "Top-10 worst: 1 case data entry error (Duplex 147% error) khó phát hiện bằng IQR đơn thuần.")

heading(doc, "5.3. Hướng phát triển", 2, MID_BLUE)
bullet(doc, "Thu thập thêm ≥ 10.000 tin từ nhiều nguồn (batdongsan.com.vn, alonhadat.com.vn, mogi.vn), mở rộng 24 quận TP.HCM.")
bullet(doc, "Feature engineering: log(area_m2), area × project_name, bedroom_density, is_high_end_district, khoảng cách đến trung tâm Quận 1 (từ lat/lon + OpenStreetMap).")
bullet(doc, "Mô hình: thử XGBoost / LightGBM với hyperparameter tuning (Optuna), thêm feature missing_count để model biết uncertainty.")
bullet(doc, "Hệ gợi ý: mở rộng thành 'quận tương đương' (cùng tier giá), thêm trọng số tuỳ biến cho budget/area, thử collaborative filtering nếu có dữ liệu hành vi người dùng.")
bullet(doc, "Text mining: trích xuất thông tin tầng/view từ title và description bằng TF-IDF hoặc LLM, bổ sung features quan trọng mà data thiếu.")
bullet(doc, "Triển khai: đóng gói thành REST API (Flask/FastAPI), xây dashboard trực quan (Streamlit).")

doc.add_page_break()

# =============================================================================
# TÀI LIỆU THAM KHẢO
# =============================================================================
heading(doc, "TÀI LIỆU THAM KHẢO", 1, DARK_BLUE)

refs = [
    "[1] Pedregosa, F., Varoquaux, G., Gramfort, A., et al. (2011). Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research, 12, 2825–2830. https://jmlr.csail.mit.edu/papers/v12/pedregosa11a.html",
    "[2] McKinney, W. (2010). Data Structures for Statistical Computing in Python. Proceedings of the 9th Python in Science Conference (SciPy 2010), 56–61. https://conference.scipy.org/proceedings/scipy2010/pdfs/Mckinney.pdf",
    "[3]scikit-learn developers. (2024). sklearn.preprocessing.StandardScaler. https://scikit-learn.org/stable/modules/generated/sklearn.preprocessing.StandardScaler.html",
    "[4] Breiman, L. (2001). Random Forests. Machine Learning, 45(1), 5–32. https://doi.org/10.1023/A:1010933404324",
    "[5] Friedman, J. H. (2001). Greedy Function Approximation: A Gradient Boosting Machine. Annals of Statistics, 29(5), 1189–1232. https://doi.org/10.1214/aos/1013203451",
    "[6] Rousseeuw, P. J. (1987). Silhouettes: A Graphical Aid to the Interpretation and Validation of Cluster Analysis. Journal of Computational and Applied Mathematics, 20, 53–65. https://doi.org/10.1016/0377-0427(87)90125-7",
    "[7] Ricci, F., Rokach, L., & Shapira, B. (2015). Recommender Systems Handbook (2nd ed.). Springer. https://doi.org/10.1007/978-1-4899-7637-6",
    "[8] James, G., Witten, D., Hastie, T., & Tibshirani, R. (2013). An Introduction to Statistical Learning (with Applications in R). Springer. https://www.statlearning.com",
    "[9] Chen, T., & Guestrin, C. (2016). XGBoost: A Scalable Tree Boosting System. Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining (KDD '16), 785–794. https://doi.org/10.1145/2939672.2939785",
    "[10] Hunter, J. D. (2007). Matplotlib: A 2D Graphics Environment. Computing in Science & Engineering, 9(3), 90–95. https://doi.org/10.1109/MCSE.2007.55",
    "[11] Đề tài Chuyên đề cuối kỳ — Môn Lập trình cho Khoa học Dữ liệu (file đề bài trong thư mục data/).",
    "[12] Dataset real_estate_apartment.xlsx — snapshot tin đăng BĐS căn hộ TP.HCM từ chotot.com (tháng 07–08/2026).",
]

for ref in refs:
    p = doc.add_paragraph()
    para_fmt(p, WD_ALIGN_PARAGRAPH.JUSTIFY, 0, 8)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    run = p.add_run(ref)
    run.font.size = Pt(11)

# =============================================================================
# SAVE
# =============================================================================
doc.save(str(OUT))
print(f"✅ Đã lưu: {OUT}")
print(f"   Kích thước: {OUT.stat().st_size / 1024:.1f} KB")
